import streamlit as st, os, re, pandas as pd, uuid, base64
import time, random
from google import genai
from google.genai import types, errors
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from streamlit_mic_recorder import mic_recorder
from concurrent.futures import ThreadPoolExecutor, as_completed

# --- KOMPLETNA STRUKTURA I UNIFIKACJA NAZEWNICTWA 1:1 Z DOKUMENTEM ANI ---
STRUKTURA_PROTOKOLU = [
    "Powód konsultacji:", "Aktualne samopoczucie:", "Aktywność:", "Apetyt:", "Pragnienie:",
    "Dotychczasowe żywienie:", "Smaczki i przysmaki:", "Ulubione smaki:",
    "### Kategorycznie tak:", "### Kategorycznie nie:", "### Kluczowa uwaga dot. przechowywania:",
    "Kał / Biegunka / Wymioty:", "Mocz:", "Odrobaczanie:", "Aktualne badania:", "Aktualne leki:",
    "Komentarz do wywiadu:", "Główne założenia diety:", "Co się zmieni na diecie BARF/BACF:",
    "Plan dietetyczny:", "Tranzycja i przechowywanie:", "Kaloryczność:", "Piciu:",
    "### Jakiej wody używać?", "Suplementy dodatkowe:", "Wiązanie fosforu:", "Smaczki:",
    "Inne smaczki:", "Karmy komercyjne:", "Tyndalizacja:", "Wprowadzanie suplementów:",
    "Badania kontrolne:", "Załączniki:"
]

# Stałe bloki edukacyjne z właściwymi, działającymi linkami MeatPoint
TEKST_TYNDALIZACJA_STALY = (
    "Jeżeli robią Państwo dietę na dłużej niż 5-6 dni (mowa o diecie surowej) i chcą Państwo "
    "ją bezpiecznie przechowywać w słoiczkach w lodówce (bez zamrażania) LUB przygotowują Państwo "
    "dietę gotowaną (BACF) na zapas, konieczne jest przeprowadzenie procesu tyndalizacji (potrójnej pasteryzacji).\n\n"
    "Proces ten skutecznie eliminuje formy przetrwalnikowe bakterii (m.in. Clostridium botulinum - jadu kiełbasianego), "
    "które mogłyby namnażać się w warunkach beztlenowych zamkniętego słoika.\n\n"
    "Pełną instrukcję krok po kroku, jak prawidłowo i bezpiecznie przeprowadzić ten proces w domowych warunkach, "
    "znajdą Państwo w [naszym artykule na blogu o tyndalizacji](https://meatpoint.io/pl/barf-wiedza/tyndalizacja-czyli-jak-przechowywac-posilki-jesli-nie-chcemy-ich-mrozic).\n\n"
    "Dodatkowo przygotowaliśmy dla Państwa [praktyczny poradnik wideo na YouTube](https://www.youtube.com/watch?v=tyfT3kmq3ME), "
    "gdzie pokazujemy cały proces krok po kroku."
)

TEKST_INNE_SMACZKI_STALY = (
    "Wprowadzając do codziennej rutyny jakiekolwiek inne smaczki komercyjne, należy bezwzględnie "
    "pamiętać o kontrolowaniu ich kaloryczności, aby nie zaburzyć bilansu nowej diety pacjenta.\n\n"
    "Szczegółowy poradnik oraz instrukcję, jak samodzielnie wyliczyć kaloryczność dowolnego produktu komercyjnego "
    "na podstawie danych z etykiety, znajdą Państwo w [naszym artykule o kaloryczności smaczków](https://meatpoint.io/pl/barf-wiedza/smaczki-i-dodatkowe-kalorie-obliczanie-kalorycznosci-komercyjnych-produktow)."
)

TEKST_WPROWADZANIE_SUPLEMENTOW_STALY = (
    "Proszę zacząć od:\n• Wody\n• Mięsa\n• Podrobów\n• tłuszczu\n• Tauryny\n"
    "Proszę przygotować dietę tylko z ich zawartością i na razie pominąć pozostałe suplementy.\n\n"
    "Jak Kicia będzie się dobrze czuła, na następny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Witaminy E\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Witaminy E\n• Witamin B\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Witaminy E\n• Witamin B\n• Jodu\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Witaminy E\n• Witamin B\n• Jodu\n• Dodatkowo: \n\n"
    "To będzie już kompletna dieta."
)


# ==============================================================================
# 📝 STYL PISANIA ANI — reguły + przykłady "draft → finał"
# ------------------------------------------------------------------------------
# To jest SERCE jakości opisów. Aby nauczyć narzędzie lepiej pisać jak Ania:
# dokładaj kolejne pary do listy PRZYKLADY_STYLU (surowy draft -> poprawka Ani).
# Im więcej konkretnych par, tym trafniej narzędzie pisze od razu w jej stylu.
# ==============================================================================

REGULY_STYLU_ANI = (
    "KIM JESTEŚ (NAJWAŻNIEJSZE):\n"
    "Jesteś sekretarzem medycznym dietetyk Anny Michalskiej. Twoim zadaniem jest WIERNIE PRZEŁOŻYĆ "
    "materiały z wizyty na profesjonalny opis w JEJ stylu. NIE jesteś drugim dietetykiem. "
    "NIE tworzysz własnych diagnoz, ocen, zaleceń ani wniosków. Referujesz to, co wynika z materiałów — nie dodajesz nic od siebie.\n\n"
    "🎯 TWÓJ GŁÓWNY CEL — PRZECZYTAJ ZANIM ZACZNIESZ:\n"
    "Twoim NADRZĘDNYM zadaniem jest WYDOBYĆ JAK NAJWIĘCEJ FAKTÓW z materiałów i uporządkować je w sekcjach. "
    "Po to powstało to narzędzie — żeby Ania nie musiała przepisywać wywiadu ręcznie.\n"
    "Poniższe zakazy dotyczą WYŁĄCZNIE interpretacji, diagnoz i zaleceń — NIE faktów.\n"
    "Bądź MAKSYMALNIE dokładny i skrupulatny w wyławianiu faktów (co powiedział Opiekun, co powiedziała Ania, "
    "co wynika z notatek). Wypełniaj sekcje treścią wszędzie tam, gdzie materiały cokolwiek na dany temat zawierają.\n"
    "Znaczniki [BRAK INFORMACJI] i [DO UZUPEŁNIENIA] to OSTATECZNOŚĆ, a nie wygodna wymówka. "
    "Dokument pełen znaczników = źle wykonana praca. Dokument pełen wydobytych faktów = dobrze wykonana praca.\n\n"
    "ZASADY STYLU ANI:\n\n"
    "1. TŁUMACZ JĘZYK OPIEKUNA NA PROFESJONALNY — nie przepisuj go dosłownie.\n"
    "   Potoczne, emocjonalne słowa Opiekuna zamień na neutralne, zawodowe.\n"
    "   ŹLE: 'Opiekunka odwaliła ogromną pracę'  ->  DOBRZE: 'Po znacznej pracy behawioralnej'\n\n"
    "2. REFERUJ, NIE DIAGNOZUJ — to granica bezpieczeństwa.\n"
    "   Opisuj stan faktyczny, nie stawiaj rozpoznań medycznych, których nie ma wprost w materiałach.\n"
    "   ŹLE: 'Otyłość u obu psów'  ->  DOBRZE: 'Nadwaga u obu psów'\n"
    "   Jeśli opiekun mówi 'tyją, ale nie wyglądają na grube' — to jest NADWAGA, nie otyłość.\n\n"
    "3. ZERO OCENIANIA I ZERO OSKARŻEŃ — zwłaszcza innych lekarzy.\n"
    "   Nigdy nie komentuj ani nie krytykuj leczenia prowadzonego przez weterynarzy.\n"
    "   ŹLE: 'Brak wiedzy weterynarza doprowadził do lekooporności'  ->  po prostu tego NIE pisz.\n"
    "   Decyzje medyczne oddawaj lekarzom: 'do decyzji lekarza prowadzącego', 'zgodnie z zaleceniami lekarza'.\n\n"
    "4. PISZ PROSTO I CIEPŁO — 'profesjonalny' to NIE 'napuszony'.\n"
    "   Pisz jak żywy, kompetentny człowiek do drugiego człowieka. Używaj 'proszę', tłumacz 'dlaczego' prosto.\n"
    "   ŹLE: 'Konieczne rygorystyczne cięcie kaloryczne i precyzyjne odmierzenie wielkości posiłków oddelegowaną do zadań z behawiorystą'\n"
    "   DOBRZE: 'ważne jest zapewnienie odpowiedniej kaloryczności, aby zapobiec przybieraniu na wadze i umożliwić psu schudnięcie'\n\n"
    "5. SELEKCJONUJ I SKRACAJ — gęściej, nie dłużej.\n"
    "   Nie przepisuj wszystkiego. Streszczaj, wyciągaj to, co istotne klinicznie. Surowe dane (np. pełny panel krwi) streść.\n"
    "   ŹLE: przepisanie każdej wartości morfologii z osobna.\n"
    "   DOBRZE: 'Morfologia w normie, biochemia w normie (mocznik nieznacznie podwyższony), Cynk w górnej granicy normy; Miedź wyraźny niedobór.'\n\n"
    "6. BĄDŹ PRECYZYJNY TAM, GDZIE TO MA ZNACZENIE KLINICZNE — skracanie nie znaczy rozmywania faktów.\n"
    "   Istotne konkrety zostają dokładne.\n"
    "   ŹLE: 'alergia na drób'  ->  DOBRZE: 'alergia na kurczaka' (jeśli zwierzę je kaczkę i indyka, to nie cały drób).\n\n"
    "7. SYGNALIZUJ NIEPEWNOŚĆ — 'może być', nie 'jest'.\n"
    "   Gdzie brak pewności, stawiaj hipotezę ostrożnie i kieruj do specjalisty. Nie orzekaj kategorycznie.\n"
    "   DOBRZE: 'Może to być miejscowe zapalenie gruczołów łojowych; proszę pokazać zmianę dermatologowi.'\n\n"
    "8. GRANICA FAKT vs INTERPRETACJA — wyciągaj WSZYSTKIE fakty, nie zmyślaj wniosków.\n"
    "   To są DWIE różne rzeczy i nie wolno ich mylić:\n"
    "   - FAKTY obecne w materiałach: wyciągaj ZAWSZE, nawet jeśli trzeba je pozbierać z różnych miejsc "
    "(transkrypcja + notatki + załączniki). Jeśli w notatkach jest 'kastrowany', napisz 'kastrowany' — NIE [BRAK INFORMACJI]. "
    "Jeśli opiekun podał BCS/wagę/wiek gdziekolwiek w materiałach, wyciągnij to. [BRAK INFORMACJI] wstawiaj TYLKO, gdy danej informacji naprawdę nigdzie nie ma.\n"
    "   - WNIOSKI, DIAGNOZY, HIPOTEZY, których nie ma wprost w materiałach: NIGDY nie dodawaj od siebie. "
    "Nie dopisuj przyczyn, mechanizmów ani teorii, których Ania nie wypowiedziała (np. 'mogło nałożyć się na wymianę włosa szczenięcego' — jeśli tego nie ma w materiałach, NIE pisz).\n"
    "   Zasada: bądź ODWAŻNY w wyciąganiu faktów, ale ZEROWY w wymyślaniu interpretacji.\n\n"
    "9. TO ANIA PODEJMUJE DECYZJE, NIE TY — nie deklaruj planu w jej imieniu.\n"
    "   Nie pisz 'wprowadzimy', 'celujemy', 'zastosujemy', 'zredukujemy' jakbyś sam układał dietę.\n"
    "   Zamiast tego opisuj, co będzie zrobione, w formie Ani: 'proszę dodawać...', 'wyliczę dokładną dawkę', "
    "'zostaną przygotowane', 'dawka uwzględniona w diecie'.\n"
    "   ŹLE: 'Wprowadzimy dodatkową suplementację tauryny'  ->  DOBRZE: 'Taurynę proszę dodawać do gotowanej mieszanki; dokładną dawkę wyliczę w diecie.'\n\n"
    "10. NIE DORADZAJ — ODNOTUJ I ZOSTAW MIEJSCE DLA ANI. (BARDZO WAŻNE)\n"
    "   Rozróżnij DWA rodzaje zaleceń:\n"
    "   a) ZALECENIA, KTÓRE ANIA WYPOWIEDZIAŁA (w transkrypcji lub notatkach) — zapisz je WIERNIE. To jest Twoje główne zadanie.\n"
    "   b) ZALECENIA, KTÓRYCH ANIA NIE WYPOWIEDZIAŁA — NIGDY ich nie wymyślaj. Nawet jeśli wydają się oczywiste lub nieszkodliwe.\n"
    "   ZAKAZANE zwroty, jeśli Ania ich nie powiedziała: 'proszę obserwować', 'warto rozważyć', 'wskazana jest konsultacja', "
    "'zaleca się', 'należy skontrolować', 'sugeruję'.\n"
    "   ZAMIAST TEGO: odnotuj sam FAKT, że Opiekun coś zgłosił, i postaw znacznik [DO UZUPEŁNIENIA] — Ania sama wpisze ocenę i zalecenie.\n"
    "   ŹLE: 'U Kafki pojawiają się tłuste zmiany na sierści. Może to być zapalenie gruczołów łojowych; proszę obserwować i pokazać dermatologowi.'\n"
    "   DOBRZE: 'Opiekunka zgłasza okresowo pojawiające się tłuste, okrągłe zmiany na sierści u Kafki, bez świądu; skóra różowa i elastyczna. [DO UZUPEŁNIENIA]'\n\n"
    "ZNACZNIKI — używaj ich świadomie, to dwie RÓŻNE rzeczy:\n"
    "   [BRAK INFORMACJI] = nikt o tym nie mówił, brak danych w materiałach.\n"
    "   [DO UZUPEŁNIENIA] = temat pojawił się w wywiadzie, fakt jest odnotowany, ale ocena/zalecenie należy do Ani.\n"
    "   Nie wstawiaj [DO UZUPEŁNIENIA] wszędzie hurtowo — tylko tam, gdzie realnie brakuje oceny Ani do zgłoszonego tematu.\n\n"
    "11. UŻYWAJ IMIENIA PACJENTA — nie 'kot', 'pies', 'pacjent'.\n"
    "   Wszędzie, gdzie brzmi to naturalnie, pisz imię zwierzęcia ('Tobiasz chętnie zjada...', 'Kicia odmawia...').\n"
    "   Określeń gatunkowych używaj tylko, gdy mowa o cechach gatunku w ogóle, nie o tym konkretnym zwierzęciu.\n"
    "   ŹLE: 'Pacjent chętnie zjada kawałki surowego kurczaka'  ->  DOBRZE: 'Tobiasz chętnie zjada kawałki surowego kurczaka'\n"
)

# Pary "draft narzędzia -> poprawka Ani". DOKŁADAJ KOLEJNE, gdy Ania je przyśle.
# Format: (surowy fragment jak napisałoby narzędzie, wersja poprawiona przez Anię).
PRZYKLADY_STYLU = [
    (
        "Opiekunka odwaliła ogromną pracę behawioralną – reaktywność w stosunku do obcych psów spadła o 90%.",
        "Po znacznej pracy behawioralnej reaktywność w stosunku do obcych psów spadła o 90%."
    ),
    (
        "Otyłość u obu psów na diecie gotowanej BARF/BACF.",
        "Nadwaga u obu psów na diecie gotowanej BACF."
    ),
    (
        "Brak wiedzy weterynarza i celowane leczenie uszu 'szerokim spektrum' w ciemno doprowadziło do lekooporności i straty czasu.",
        "W przypadku nawrotu infekcji uszu po odstawieniu maści z antybiotykiem wskazana jest konsultacja dermatologiczna celem pobrania wymazu i oceny cytologicznej."
    ),
    (
        "Konieczne rygorystyczne cięcie kaloryczne i precyzyjne odmierzenie wielkości posiłków BACF u obojga pacjentów. Posiłki muszą uwzględniać stałą pulę 10-15% dziennej dawki energetycznej w pełni oddelegowaną do zadań z behawiorystą i do oswajania bodźców medycznych.",
        "Ważne jest zapewnienie odpowiedniej kaloryczności, aby zapobiec przybieraniu na wadze i umożliwić psom schudnięcie. Uwzględniamy pulę smakołyków (10-15% dziennej dawki kalorii), której Opiekunka powinna restrykcyjnie przestrzegać — to warunek skutecznej redukcji masy ciała."
    ),
    (
        "Ze względu na starszą siostrę Kafkę, u której występuje alergia na drób, z diety Frania całkowicie wykluczono kurczaka.",
        "Ze względu na starszą siostrę Kafkę, u której występuje alergia na kurczaka, z diety Frania całkowicie wykluczono kurczaka."
    ),
    (
        "Sterylizacja/kastracja: [BRAK INFORMACJI]",
        "Sterylizacja/kastracja: kastrowany  (fakt był w notatkach z wizyty — należało go wyciągnąć, a nie wstawiać [BRAK INFORMACJI])"
    ),
    (
        "Wprowadzimy dodatkową suplementację tauryny (konieczną przy obróbce termicznej) oraz leczniczą dawkę magnezu. U Kafki celujemy w około 370-380 kcal na dzień.",
        "Taurynę proszę dodawać do gotowanej mieszanki zawsze (wyliczę dokładną dawkę, ponieważ są to diety gotowane). Magnez w wyliczonej dawce poza standardowym przepisem. U Kafki proszę celować w kaloryczność ok. 370-380 kcal na dzień."
    ),
    (
        "Zmiana struktury włosa może mieć związek z niedoborem miedzi, jednak mogła również nałożyć się na naturalny okres wymiany włosa szczenięcego na dorosły w czasie, gdy pies przebywał na diecie komercyjnej.",
        "Niedobór miedzi (przy podwyższonym cynku) jest odpowiedzialny za odbarwienie i wełnianą strukturę włosa. Suplementacja wyrównująca poprawi wybarwienie u nasady, jednak w pełni uformowany włos może pozostać zmieniony aż do wymiany w procesie linienia. (Pisz tylko to, co wynika z materiałów Ani — nie dodawaj własnych teorii o okresie dorastania, jeśli ich tam nie ma.)"
    ),
    (
        "U Kafki okazjonalnie pojawiają się tłuste, okrągłe zmiany na sierści. Może to być miejscowe zapalenie gruczołów łojowych — proszę zmianę obserwować i przy nawrocie pokazać ją dermatologowi.",
        "Opiekunka zgłasza okresowo pojawiające się u Kafki tłuste, okrągłe zmiany na sierści, bez świądu; skóra różowa i elastyczna. [DO UZUPEŁNIENIA]"
    ),
    (
        "Aktualne badania: Morfologia (WBC 13.52 tys., RBC 7.02 mln, HGB 16 g%) - w normie. Kreatynina 0.9 mg/dl - w normie. Mikroelementy: Cynk 199.0 µg/dl - w górnej granicy normy; Miedź 65.8 µg/dl – wyraźny i bardzo silny niedobór.",
        "Aktualne badania: Badanie krwi Frania z dn. 25.06.2026 r. — morfologia i biochemia w normie, miedź wyraźny niedobór. (Zapisuj TYLKO to, co Ania sama powiedziała o badaniach — nie odczytuj i nie przepisuj wartości z załączników ani zdjęć.)"
    ),
    (
        "Zaleca się ograniczenie podawania suplementów bezpośrednio do jamy ustnej. Priorytetem jest akceptacja diety przez pacjenta.",
        "Liczy się dobrostan kota. Kicia karmiona paręnaście razy dziennie suplementami zamiast dietą nie jest tym, co można nazwać kotem w dobrostanie. Zależy nam przede wszystkim, żeby akceptowała dietę. Kot musi mieć przyjemność z jedzenia i z życia samego w sobie."
    ),
    (
        "Opiekunka błędnie zakłada, że karmy nerkowe będą odpowiednie dla pacjentki. Jest to nieprawidłowy wybór.",
        "Kicia to nie jest po prostu „kot nerkowy”. To kot z zaburzeniami wchłaniania żelaza, chorobą jelit, przewlekłym zapaleniem trzustki, euthyroid sick syndrome itp. Karmy nerkowe charakteryzują się bardzo dużą ilością tłuszczu i węglowodanów, a skrajnie niską ilością białka — w przypadku Kici restrykcyjne karmy nerkowe nie są wskazane."
    ),
    (
        "Kategorycznie nie: ziemniaki, cukinia ze skórką, zbyt czerwone gatunki mięsa.",
        "Kategorycznie nie:\n• (Prawdopodobnie) Ziemniaki – zwłaszcza po rozmrożeniu.\n• Cukinia podawana ze skórką lub nasionami.\n• (Prawdopodobnie) Zbyt czerwone i intensywne gatunki mięsa (np. wyrazista karkówka wieprzowa) – choć był moment, kiedy chętnie zjadła ją jako mięso surowe.\n(Ania oznacza niepewność słowem „(Prawdopodobnie)” i ZACHOWUJE sprzeczne obserwacje zamiast je wygładzać.)"
    ),
    (
        "Pacjent chętnie zjada kawałki surowego kurczaka lub wątróbkę delikatnie podsmażoną na suchej patelni. Kot wykazuje ogólne dobre zainteresowanie surowym mięsem.",
        "Tobiasz chętnie zjada kawałki surowego kurczaka lub wątróbkę delikatnie podsmażoną na suchej patelni. Wykazuje ogólne dobre zainteresowanie surowym mięsem. (Używaj imienia pacjenta zamiast 'pacjent'/'kot'/'pies'.)"
    ),
    (
        "Komentarz do wywiadu: [DO UZUPEŁNIENIA]",
        "Komentarz do wywiadu:\n- Tobiasz ma specyficzne potrzeby żywieniowe – musi spożywać mokrą karmę, ale nie jest do niej przekonany, przez co cały czas je karmę suchą.\n- Dla zdrowia jego dróg moczowych musimy przestawić go na dietę mokrą, a skoro już wykazuje zainteresowanie surowym mięsem, jesteśmy na dobrej drodze, żeby sprawnie przekonać go do diety BARF.\n(Ta sekcja to SYNTEZA wywiadu — streszczaj problemy Opiekuna i cele Ani wypowiedziane w materiałach. NIE zostawiaj jej pustej.)"
    ),
    (
        "Dieta opiera się na suchej karmie Orijen Regional Red (podawane łącznie około 70 g dziennie). Próby wprowadzenia komercyjnych karm mokrych zakończyły się niepowodzeniem. Przed adopcją kot jadł głównie suchą karmę.",
        "- Dieta opiera się na suchej karmie Orijen Regional Red (ok. 70 g dziennie, w 2-3 posiłkach).\n- Próby wprowadzenia komercyjnych karm mokrych zakończyły się niepowodzeniem.\n- Przed adopcją, w fundacji, Tobiasz również jadł głównie suchą karmę.\n(Sekcje z wieloma niezależnymi faktami formatuj jako listę punktów '- ' — tak pisze Ania.)"
    ),
]


def zbuduj_instrukcje_stylu():
    """Składa reguły stylu + przykłady draft->finał w jeden blok do promptu."""
    blok = REGULY_STYLU_ANI + "\n\nPRZYKŁADY — jak NARZĘDZIE napisało (ŹLE) i jak poprawiła to Ania (DOBRZE):\n"
    for i, (zle, dobrze) in enumerate(PRZYKLADY_STYLU, 1):
        blok += f"\nPrzykład {i}:\n  ŹLE (draft): {zle}\n  DOBRZE (Ania): {dobrze}\n"
    return blok


def czytelna_nazwa(s):
    """Wersja nazwy sekcji do POKAZANIA użytkownikowi — bez znaczników markdown (###/##)."""
    return s.replace("### ", "").replace("###", "").replace("## ", "").replace("##", "").strip()


def _norm_naglowek(s):
    """Normalizuje etykietę nagłówka do porównania: usuwa znaczniki markdown,
    bierze część do pierwszego dwukropka i podnosi do wielkich liter."""
    s = s.replace("###", "").replace("##", "").strip()
    if ":" in s:
        s = s.split(":", 1)[0]
    return s.strip().upper()


def segmentuj_docx(file_bytes):
    doc = Document(BytesIO(file_bytes))
    # Mapa: znormalizowana etykieta -> oryginalny klucz z STRUKTURA_PROTOKOLU
    naglowki = {_norm_naglowek(n): n for n in STRUKTURA_PROTOKOLU}
    sekcje = {}
    biezaca = "Nagłówek i Metryczka"
    sekcje[biezaca] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        klucz = _norm_naglowek(t)
        # Dokładne dopasowanie etykiety (nie podłańcuch) — eliminuje kolizję
        # "Smaczki:" ⊂ "Inne smaczki:" i podobne.
        if klucz in naglowki and len(t) < 65:
            biezaca = naglowki[klucz]
            sekcje[biezaca] = []
        else:
            sekcje[biezaca].append(t)
    return {k: "\n".join(v) for k, v in sekcje.items()}


def add_hyperlink(p, url, text):
    part = p.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), r_id)
    nr = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    nr.append(rPr); tn = OxmlElement('w:t'); tn.text = text; nr.append(tn); hl.append(nr); p._p.append(hl)
    return hl

# Znaczniki wymagające uwagi Ani — renderowane kolorem i pogrubieniem w Wordzie,
# żeby były nie do przeoczenia przy przeglądaniu dokumentu.
ZNACZNIKI_UWAGI = {
    '[BRAK INFORMACJI]': RGBColor(220, 38, 38),    # czerwony  — nikt o tym nie mówił
    '[DO UZUPEŁNIENIA]': RGBColor(217, 119, 6),    # pomarańczowy — czeka na ocenę Ani
}


# Zapis linku w stylu markdown: [widoczna etykieta](https://adres) — renderowany
# w Wordzie jako klikalny link ukryty pod etykietą (bez brzydkiego gołego URL-a).
WZOR_MD_LINK = re.compile(r'\[([^\]]+)\]\((https?://[^\s)]+)\)')


def _dodaj_zwykly_tekst(p, tekst):
    """Dodaje zwykły tekst do akapitu z obsługą pogrubień (**) i surowych URL-i."""
    sub_segs = tekst.split('**')
    for idx, sub_seg in enumerate(sub_segs):
        if not sub_seg:
            continue
        czy_pogrubiony = (idx % 2 == 1)
        url_segs = re.split(r'(https?://[^\s]+)', sub_seg)
        for u_idx, u_seg in enumerate(url_segs):
            if u_idx % 2 == 1:
                add_hyperlink(p, u_seg, u_seg)
            else:
                run = p.add_run(u_seg)
                if czy_pogrubiony:
                    run.bold = True


def _dodaj_tekst_z_formatowaniem(p, tekst):
    """Dodaje tekst do akapitu: najpierw linki markdown [etykieta](url),
    potem pogrubienia (**) i surowe adresy jako fallback."""
    czesci = WZOR_MD_LINK.split(tekst)
    # split z 2 grupami daje sekwencję: [zwykły, etykieta, url, zwykły, etykieta, url, ...]
    i = 0
    while i < len(czesci):
        if i % 3 == 0:
            if czesci[i]:
                _dodaj_zwykly_tekst(p, czesci[i])
            i += 1
        else:
            etykieta = czesci[i].strip()
            url = czesci[i + 1]
            add_hyperlink(p, url, etykieta if etykieta else url)
            i += 2


def parsuj_i_formatuj_tekst(p, tekst):
    """Renderuje tekst, wyróżniając kolorem znaczniki [BRAK INFORMACJI] i [DO UZUPEŁNIENIA]."""
    wzor = r'(\[BRAK INFORMACJI\]|\[DO UZUPEŁNIENIA\])'
    for seg in re.split(wzor, tekst):
        if not seg:
            continue
        if seg in ZNACZNIKI_UWAGI:
            ra = p.add_run(seg)
            ra.bold = True
            ra.font.color.rgb = ZNACZNIKI_UWAGI[seg]
        else:
            _dodaj_tekst_z_formatowaniem(p, seg)

def konwertuj_do_docx(tekst_md):
    doc = Document()
    for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin, s.header_distance = Inches(1.3), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.4)
    style = doc.styles['Normal']; font = style.font; font.name, font.size, style.paragraph_format.line_spacing, style.paragraph_format.space_after = 'Arial', Pt(10.5), 1.25, Pt(4)
    sec = doc.sections[0]; sec.different_first_page_header_footer = True
    
    t_h = sec.first_page_header.add_table(1, 2, Inches(6.7)); t_h.autofit = False
    t_h._tbl.tblPr.append(OxmlElement('w:tblBorders'))
    kl, kp = t_h.rows[0].cells[0], t_h.rows[0].cells[1]; kl.width, kp.width = Inches(4.9), Inches(1.8)
    
    pk = kl.paragraphs[0]; pk.paragraph_format.space_after = Pt(0)
    pk.add_run("Anna Michalska\n").bold = True; pk.runs[-1].font.size = Pt(11)
    ps = pk.add_run("Dietetyka Psów i Kotów\n"); ps.font.size, ps.font.color.rgb = Pt(9), RGBColor(100, 116, 139)
    pd_t = pk.add_run("miesnepsokotki@gmail.com  |  https://www.facebook.com/meatpoint.io"); pd_t.font.size, pd_t.font.color.rgb = Pt(8.5), RGBColor(100, 116, 139)
    
    if os.path.exists("logo.png"):
        kp.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        kp.paragraphs[0].paragraph_format.space_after = Pt(0)
        kp.paragraphs[0].add_run().add_picture("logo.png", width=Inches(1.0))
        sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sec.header.paragraphs[0].paragraph_format.space_after = Pt(0)
        sec.header.paragraphs[0].add_run().add_picture("logo.png", width=Inches(1.0))

    for line in tekst_md.split('\n'):
        l_s = line.strip()
        if not l_s: continue
        
        if "DATA WIZYTY:" in l_s.upper():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(12), Pt(6)
            poczatek_daty, *koniec_daty = l_s.split(':', 1)
            poczatek_czysty = poczatek_daty.replace('**', '').strip()
            p.add_run(poczatek_czysty + ': ').bold = True
            if koniec_daty: parsuj_i_formatuj_tekst(p, koniec_daty[0].strip())
            continue

        if l_s.startswith('## '):
            p = doc.add_paragraph(); p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(14), Pt(4)
            czysty_h2 = l_s.replace('## ', '').replace('**', '')
            r = p.add_run(czysty_h2); r.bold = True; r.font.size, r.font.color.rgb = Pt(12), RGBColor(194, 65, 12)
        elif l_s.startswith('### '):
            p = doc.add_paragraph(); p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(8), Pt(2)
            czysty_h3 = l_s.replace('### ', '').replace('**', '')
            r = p.add_run(czysty_h3); r.bold, r.font.size = True, Pt(10.5)
        elif l_s.startswith('- ') or l_s.startswith('* ') or l_s.startswith('• '):
            c_t = l_s.lstrip('-*• ').strip()
            p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
            if ':' in c_t and not c_t.strip().startswith('http'):
                pk_s, zk_s = c_t.split(':', 1)
                if len(pk_s) < 45: 
                    pk_czysty = pk_s.replace('**', '').strip()
                    p.add_run(pk_czysty + ': ').bold = True
                    parsuj_i_formatuj_tekst(p, zk_s)
                    continue
            parsuj_i_formatuj_tekst(p, c_t)
        else:
            if ':' in l_s and not l_s.strip().startswith('http'):
                pk_s, zk_s = l_s.split(':', 1)
                if len(pk_s) < 45: 
                    p = doc.add_paragraph()
                    pk_czysty = pk_s.replace('**', '').strip()
                    p.add_run(pk_czysty + ': ').bold = True
                    parsuj_i_formatuj_tekst(p, zk_s)
                    continue
            p = doc.add_paragraph(); parsuj_i_formatuj_tekst(p, l_s)
            
    b = BytesIO(); doc.save(b); return b.getvalue()


def generuj_z_retry(client, model_name, contents, config=None, max_prob=5):
    """Wywołuje generate_content z ponawianiem przy błędach limitu (429/quota) i
    przejściowych błędach serwera (5xx). Backoff wykładniczy + jitter.
    Inne błędy rzuca natychmiast."""
    for proba in range(max_prob):
        try:
            return client.models.generate_content(
                model=model_name, contents=contents, config=config
            )
        except errors.APIError as e:
            kod = getattr(e, "code", None)
            tekst = str(e).lower()
            czy_limit = kod == 429 or any(x in tekst for x in ["quota", "rate", "resource_exhausted"])
            czy_serwer = isinstance(kod, int) and kod >= 500
            if (czy_limit or czy_serwer) and proba < max_prob - 1:
                time.sleep((2 ** proba) + random.uniform(0, 1))
                continue
            raise


def czytelny_blad(e, kontekst=""):
    """Zamienia techniczny błąd API na zrozumiały komunikat po polsku z instrukcją,
    co użytkownik ma zrobić. Nieznane błędy zwraca w oryginale (dla diagnostyki)."""
    tekst = str(e).lower()
    kod = getattr(e, "code", None)

    # 503 / przeciążenie serwerów Google (najczęstsze przy trybie Dokładnym - model preview)
    if kod == 503 or "unavailable" in tekst or "high demand" in tekst or "overloaded" in tekst:
        return (
            "⏳ **Serwery Google są chwilowo przeciążone** (zbyt wielu użytkowników naraz).\n\n"
            "To **nie jest błąd aplikacji** — mija samo, zwykle w kilka–kilkanaście minut.\n\n"
            "**Co zrobić:**\n"
            "1. Odczekaj chwilę i kliknij przycisk ponownie (transkrypcja i załączniki zostają na miejscu), **albo**\n"
            "2. Przełącz w panelu bocznym tryb AI na **💨 Standardowy** — działa na stabilniejszym modelu "
            "i praktycznie nie łapie przeciążeń.\n\n"
            "💡 Tryb *Dokładny* najczęściej bywa przeciążony wieczorami."
        )

    # 429 / wyczerpany limit zapytań
    if kod == 429 or "quota" in tekst or "resource_exhausted" in tekst or "rate limit" in tekst:
        return (
            "🚦 **Przekroczony limit zapytań do Google** (zbyt wiele operacji w krótkim czasie).\n\n"
            "**Co zrobić:** odczekaj 1–2 minuty i spróbuj ponownie. "
            "Jeśli powtarza się często — sprawdź limity swojego konta Google AI Studio."
        )

    # 404 / nieistniejący lub wycofany model
    if kod == 404 or "not found" in tekst:
        return (
            "🔍 **Wybrany model AI jest niedostępny** (mógł zostać wycofany przez Google).\n\n"
            "**Co zrobić:** przełącz tryb AI w panelu bocznym na inny. "
            "Listę aktualnie działających modeli zobaczysz w panelu bocznym w sekcji "
            "*🔍 Pokaż dostępne modele (diagnostyka)*."
        )

    # 400/401/403 - problem z kluczem API
    if kod in (400, 401, 403) or "api key" in tekst or "permission" in tekst or "unauthenticated" in tekst:
        return (
            "🔑 **Problem z kluczem API.**\n\n"
            "**Co zrobić:** sprawdź, czy klucz w panelu bocznym jest poprawny i aktywny "
            "oraz czy na koncie Google AI Studio jest włączone rozliczanie (płatny plan)."
        )

    # Problem z pobraniem bazy artykułów z Arkusza Google
    if "arkusz" in tekst or "columns" in tekst or "csv" in tekst:
        return f"📊 **Problem z bazą artykułów (Arkusz Google):**\n\n{e}"

    # Nieznany błąd - pokaż oryginał, żeby dało się zdiagnozować
    return f"⚠️ Wystąpił nieoczekiwany błąd{kontekst}:\n\n`{e}`"


@st.cache_data(ttl=3600)
def pobierz_baze_artykulow(url):
    """Pobiera bazę artykułów z Arkusza Google (cache 1h) i waliduje kolumny."""
    csv_url = url.replace('/edit?usp=sharing', '/export?format=csv')
    df = pd.read_csv(csv_url)
    brakujace = {'URL', 'Nazwa', 'Opis dla AI'} - set(df.columns)
    if brakujace:
        raise ValueError(
            f"Arkusz Google nie zawiera wymaganych kolumn: {brakujace}. "
            f"Sprawdź nagłówki (dokładna pisownia: 'URL', 'Nazwa', 'Opis dla AI')."
        )
    return df


def przetworz_jedno_nagranie(args):
    """
    Przetwarza jedno nagranie audio: transkrypcja + edycja sekcji.
    Zwraca (nazwa_sekcji, nowa_treść) lub rzuca wyjątek.
    """
    s_nazwa, a_bytes, oryginalna_tresc, api_key, model_choice = args

    # Własny klient w każdym wątku — bezpieczne przy równoległości (bez globalnego stanu)
    client = genai.Client(api_key=api_key)

    # Krok 1: transkrypcja audio (zawsze flash — szybki i tani)
    audio_part = types.Part.from_bytes(data=a_bytes, mime_type="audio/wav")
    p_trans = "Przetwórz to nagranie audio i zwróć dokładny tekst (transkrypcję) tego, co zostało powiedziane, słowo w słowo, po polsku."
    transkrypcja_uwagi = generuj_z_retry(client, "gemini-3.5-flash", [p_trans, audio_part]).text.strip()

    # Krok 2: integracja uwagi z oryginalną treścią sekcji
    p_ed = (
        f"Jesteś precyzyjnym edytorem dokumentacji medycznej zwierząt BARF/BACF.\n"
        f"Zaktualizuj oryginalny tekst sekcji '{s_nazwa}' wyłącznie o fakty podyktowane w uwadze głosowej.\n\n"
        f"Oryginalna treść sekcji:\n{oryginalna_tresc}\n\n"
        f"Podyktowana uwaga głosowa:\n{transkrypcja_uwagi}\n\n"
        f"ZASADY:\n"
        f"1. Dołącz NOWE fakty z uwagi głosowej, zachowując spójność z oryginałem.\n"
        f"2. ZAKAZ dodawania własnych zaleceń, komentarzy AI ani wniosków.\n"
        f"3. ZAKAZ wstępów i podsumowań ('Oto zaktualizowana sekcja:' itp.).\n"
        f"4. Zwróć WYŁĄCZNIE czystą, zaktualizowaną treść sekcji."
    )
    nowa_tresc = generuj_z_retry(client, model_choice, p_ed).text.strip()
    return s_nazwa, nowa_tresc


st.set_page_config(page_title="MeatPoint - Asystent Dietetyka", layout="wide", page_icon="🐾")

with st.sidebar:
    st.header("🔑 Autoryzacja")
    api_key = st.text_input("Klucz API Gemini", type="password")
    # Przyjazne nazwy widoczne dla użytkownika -> techniczne ID modelu Google
    MODELE = {
        "💨 Standardowy — szybki": "gemini-3.5-flash",
        "🎯 Dokładny — precyzyjny": "gemini-3.1-pro-preview",
    }
    etykieta_modelu = st.selectbox("Wybierz tryb AI", list(MODELE.keys()))
    model_choice = MODELE[etykieta_modelu]
    st.caption(
        "💡 **Standardowy** — do większości wizyt, najszybszy.\n\n"
        "**Dokładny** — gdy dołączasz skomplikowane wyniki badań (PDF / zdjęcia) "
        "i zależy Ci na maksymalnej precyzji. Wolniejszy i droższy."
    )

    with st.expander("🔍 Pokaż dostępne modele (diagnostyka)"):
        if not api_key:
            st.caption("Wpisz najpierw klucz API powyżej, aby zobaczyć listę.")
        else:
            try:
                client_diag = genai.Client(api_key=api_key)
                dostepne = [
                    mm.name.replace("models/", "")
                    for mm in client_diag.models.list()
                    if mm.supported_actions and "generateContent" in mm.supported_actions
                ]
                st.caption("Modele wspierające generateContent na Twoim koncie:")
                st.code("\n".join(sorted(dostepne)) or "brak")
            except Exception as e:
                st.caption(f"Nie udało się pobrać listy: {e}")

    st.markdown("---")
    st.info(
        "🔒 **Bezpieczeństwo danych pacjenta:**\n\n"
        "Narzędzie przetwarza dane w bezpiecznym, szyfrowanym strumieniu bezpośrednio przez oficjalne Google Gemini API (płatny plan).\n\n"
        "- Transkrypcje i załączniki **NIE** są zapisywane na serwerach.\n\n"
        "- Na płatnym planie dane **NIE** są wykorzystywane do trenowania modeli AI.\n\n"
        "- Po zamknięciu karty przeglądarki cała sesja bezpowrotnie znika z pamięci."
    )

tab1, tab2 = st.tabs(["🚀 Generator opisów wizyt (Wklej Tekst)", "🎙️ Edytor głosowy opisów wizyt"])

# ==============================================================================
# 🚀 ZAKŁADKA 1
# ==============================================================================
with tab1:
    st.title("🐾 Generator opisów wizyt")
    st.markdown("Wklej transkrypcję i dołącz dowolne załączniki (notatki, stare plany, wyniki badań w PDF, zdjęciach lub plikach Word .docx). AI inteligentnie przyporządkuje informacje do odpowiednich sekcji.")
    
    col1, col2 = st.columns([1, 1], gap="large")
    
    with col1:
        transcript = st.text_area("🔊 Wklej tutaj kompletną transkrypcję z rozmowy:", height=380, key="surowy_wklejony_tekst")
        
        zalaczniki = st.file_uploader(
            "📂 Dołącz załączniki (PDF, Word .docx, notatki, zdjęcia dokumentacji itp.):", 
            type=["pdf", "png", "jpg", "jpeg", "docx"], 
            accept_multiple_files=True,
            key="pliki_kliniczne"
        )
        if zalaczniki:
            st.success(f"📎 Pomyślnie załadowano załączniki kontekstowe: {len(zalaczniki)} szt.")
        
    with col2:
        st.subheader("📋 Opis wizyty")
        LINK_DO_ARKUSZA = "https://docs.google.com/spreadsheets/d/1qgSX_t4_fb36CqtFUluPDKDQILpR9_SLOlYBPTXSTes/edit?usp=sharing"
        
        if st.button("🚀 Wygeneruj i uzupełnij opisy wizyty w Word", type="primary", use_container_width=True):
            if not api_key or not transcript: 
                st.error("❌ Uzupełnij klucz API oraz upewnij się, że okno transkrypcji nie jest puste!")
            else:
                with st.spinner("Globalna analiza transkrypcji i załączników oraz dopasowywanie nagłówków..."):
                    try:
                        df = pobierz_baze_artykulow(LINK_DO_ARKUSZA); l_p = ""
                        for _, r in df.iterrows(): l_p += f"- Link: {r['URL']} | Tytuł: {r['Nazwa']} | Kiedy dołączyć (Wskazanie): {r['Opis dla AI']}\n"
                        
                        client = genai.Client(api_key=api_key)
                        config_gen = types.GenerateContentConfig(
                            system_instruction=(
                                zbuduj_instrukcje_stylu()
                                + "\n\n"
                                + "TWOJE ZADANIE TECHNICZNE:\n"
                                "Stwórz jeden, spójny opis wizyty na podstawie trzech źródeł: ustnej transkrypcji, "
                                "przesłanych dokumentów/zdjęć (załączników) oraz tekstów z plików Word. "
                                "Stosuj wszystkie powyższe zasady stylu Ani w KAŻDYM zdaniu, które piszesz.\n\n"
                                "ZASADA INTELIGENTNEGO DOPASOWANIA (CROSS-ANALYSIS):\n"
                                "1. Przeanalizuj treść każdego załącznika. Informacje w nich zawarte mogą dotyczyć DOWOLNEJ sekcji protokołu (notatki o wodzie, uwagi o smaczkach, dawki leków, opisy samopoczucia, wyniki badań).\n"
                                "2. Przyporządkuj fakty tematycznie: informacje o diecie komercyjnej do 'Karmy komercyjne', informacje o dawkowaniu wody do 'Piciu/Jakiej wody używać', a opisy dolegliwości do 'Powód konsultacji' lub 'Kał/Biegunka/Wymioty'.\n"
                                "   🚨 WYJĄTEK — WYNIKI BADAŃ LABORATORYJNYCH: jeśli w załącznikach są surowe wyniki badań (PDF/zdjęcie z laboratorium, tabele z wartościami) — NIE odczytuj ich, NIE przepisuj wartości i NIE interpretuj. "
                                "Do sekcji 'Aktualne badania' trafia WYŁĄCZNIE to, co Ania sama powiedziała lub napisała o badaniach własnymi słowami.\n"
                                "   Uwaga na rozróżnienie: jeśli w NOTATKACH ANI jest zdanie typu 'kreatynina wzrosła', 'morfologia w normie' — to są JEJ słowa i MASZ je wykorzystać. "
                                "Zakaz dotyczy wyłącznie samodzielnego odczytywania liczb z dokumentów laboratoryjnych.\n"
                                "3. Zintegruj wiedzę z transkrypcji i załączników. Jeśli dokumenty i transkrypcja mówią o tym samym, połącz te fakty w spójny opis.\n\n"
                                "ZASADY OGÓLNE:\n"
                                "- Pisz WYŁĄCZNIE prawdę na podstawie dostarczonych materiałów. ZAKAZ zmyślania faktów czy dawek.\n"
                                "- Jeśli chcesz coś wyróżnić, używaj podwójnych gwiazdek **tekst**.\n"
                                "- Jeśli w źródłach brakuje danych dla danej sekcji, wstaw fragment [BRAK INFORMACJI].\n"
                                "- Jeśli temat pojawił się w wywiadzie, ale brakuje oceny/zalecenia Ani — odnotuj sam fakt "
                                "i wstaw [DO UZUPEŁNIENIA]. NIE wymyślaj zalecenia za Anię.\n"
                                "- KAŻDY link wstawiaj w formacie [krótka etykieta](URL) — nigdy goły adres URL.\n"
                                "- Sekcje zawierające kilka niezależnych faktów formatuj jako listę punktów zaczynających się od '- ' "
                                "(tak pisze Ania), zamiast zbijać je w jeden akapit prozy."
                            )
                        )
                        
                        instrukcja_szablonu = ""
                        for naglowek in STRUKTURA_PROTOKOLU:
                            if naglowek == "Załączniki:":
                                instrukcja_szablonu += f"## {naglowek}\n- Dołącz wyłącznie pasujące linki z bazy, jeśli ich warunki kliniczne zostały spełnione. Każdy link wstaw w formacie [Nazwa artykułu](URL).\n- Pod nimi dodaj dokładnie te słowa:\nW razie pytań dotyczących tego opisu, jestem do Państwa dyspozycji.\nZachęcamy również do poszerzenia wiedzy o diecie na [naszej stronie meatpoint.io](https://meatpoint.io) lub [na Facebooku](https://www.facebook.com/meatpoint.io)\n\nPozdrawiam serdecznie,\nAnna Michalska"
                            elif naglowek == "Tyndalizacja:":
                                instrukcja_szablonu += f"## {naglowek}\n{TEKST_TYNDALIZACJA_STALY}\n\n"
                            elif naglowek == "Inne smaczki:":
                                instrukcja_szablonu += f"## {naglowek}\n{TEKST_INNE_SMACZKI_STALY}\n\n"
                            elif naglowek == "Wprowadzanie suplementów:":
                                instrukcja_szablonu += f"## {naglowek}\n{TEKST_WPROWADZANIE_SUPLEMENTOW_STALY}\n\n"
                            elif naglowek == "Komentarz do wywiadu:":
                                instrukcja_szablonu += f"## {naglowek}\n- Napisz zwięzłą syntezę wywiadu w punktach '- ': (1) najważniejsze problemy i oczekiwania Opiekuna, (2) co Ania chce osiągnąć i dlaczego — WYŁĄCZNIE na podstawie tego, co padło w transkrypcji/notatkach.\n- To jest STRESZCZENIE materiałów, nie nowe porady — nie dodawaj zaleceń, których Ania nie wypowiedziała.\n- NIE zostawiaj tej sekcji pustej, jeśli wywiad zawiera jakiekolwiek problemy/cele. [DO UZUPEŁNIENIA] wstaw tylko w miejscu celów Ani, jeśli w materiałach ich nie wyraziła.\n"
                            elif naglowek == "Aktualne badania:":
                                instrukcja_szablonu += f"## {naglowek}\n- Zapisz WYŁĄCZNIE to, co sama Ania powiedziała lub napisała o badaniach WŁASNYMI SŁOWAMI — w transkrypcji LUB w swoich notatkach (np. 'morfologia w normie', 'miedź wyraźny niedobór', 'kreatynina wzrosła', data badania). To są jej słowa i MASZ je wykorzystać.\n- 🚨 ZAKAZ samodzielnego odczytywania i przepisywania wartości liczbowych z surowych dokumentów laboratoryjnych (PDF, zdjęcia, tabele wyników). ZAKAZ własnej oceny czy interpretacji wyników.\n- Jeśli Ania nic o badaniach nie powiedziała ani nie napisała, wstaw [BRAK INFORMACJI].\n"
                            elif naglowek == "Badania kontrolne:":
                                instrukcja_szablonu += f"## {naglowek}\n- Przedstaw WYŁĄCZNIE badania kontrolne, które Ania sama wskazała w transkrypcji lub notatkach, wraz z podanymi przez nią terminami.\n- ZAKAZ proponowania własnych badań kontrolnych. Jeśli Ania nic nie wskazała, wstaw [DO UZUPEŁNIENIA].\n"
                            else:
                                prefix = "" if naglowek.startswith("###") else "## "
                                instrukcja_szablonu += f"{prefix}{naglowek}\n- Analizuj pod kątem tego nagłówka zarówno tekst transkrypcji, jak i dołączone pliki załączników. Wyciągnij precyzyjne fakty.\n"

                        pakiety_danych_dla_ai = []
                        teksty_z_docx = ""
                        
                        if zalaczniki:
                            for plik in zalaczniki:
                                if plik.name.endswith(".docx"):
                                    doc_ctx = Document(BytesIO(plik.read()))
                                    akapit_tekst = [p.text for p in doc_ctx.paragraphs if p.text.strip()]
                                    teksty_z_docx += f"\n--- ZAWARTOŚĆ DOŁĄCZONEGO PLIKU WORD ({plik.name}) ---\n" + "\n".join(akapit_tekst) + "\n"
                                else:
                                    bytes_data = plik.read()
                                    pakiety_danych_dla_ai.append(
                                        types.Part.from_bytes(data=bytes_data, mime_type=plik.type)
                                    )
                        
                        prompt_glowny = f"Przeanalizuj podaną transkrypcję wizyty oraz wszystkie dołączone pliki kontekstowe.\n\nWygeneruj dokument według tej rygorystycznej kolejności:\n\nKROK 1: Na samej górze stwórz wyrównaną DO LEWEJ linię: 'Data wizyty: DD.MM.YYYY' (wyciągnij datę z rozmowy/plików lub wstaw [BRAK INFORMACJI])\n\nKROK 2: Bezpośrednio POD DATĄ wypisz linie metryczki podstawowej (ZAKAZ używania znaków '##' na ich początku, po dwukropku ma być dokładnie jedna spacja. Dane wyciągaj z transkrypcji oraz załączników):\nDane Opiekuna: \nPacjent: \nGatunek: \nRasa: \nWiek: \nWaga: \nBCS: \nMCS: \nIlość zwierząt w domu: \nSterylizacja/kastracja: \n\nKROK 3: Pod metryczką umieść poniższe nagłówki i uzupełnij je danymi z transkrypcji oraz plików, zachowując ich identyczną wielkość liter i pisownię:\n{instrukcja_szablonu}\n\n🚨 DEDYKOWANE DOPASOWANIE LINKÓW Z ARKUSZA:\nOto dostępna baza załączników zewnętrznych:\n{l_p}\n\nPrzeanalizuj pole 'Kiedy dołączyć (Wskazanie)'. Dołącz dany adres URL do dokumentu TYLKO wtedy, gdy z transkrypcji lub przesłanych załączników wynika, że pacjent cierpi na opisaną dolegliwość. Jeśli brak dopasowania, pomiń link. Każdy dołączany link wstaw w formacie [Tytuł artykułu](URL) — NIGDY nie wklejaj gołego adresu URL.\n\nTranskrypcja rozmowy:\n{transcript}\n"
                        
                        if teksty_z_docx:
                            prompt_glowny += f"\nDodatkowe dokumenty tekstowe przesłane w załącznikach Word:\n{teksty_z_docx}"
                        
                        pakiety_danych_dla_ai.append(prompt_glowny)
                        
                        res = generuj_z_retry(client, model_choice, pakiety_danych_dla_ai, config=config_gen)
                        
                        st.text_area("Podgląd tekstu wynikowego:", value=res.text, height=350, key="podglad_gen")
                        st.download_button("📥 POBIERZ GOTOWY PLIK WORD (.DOCX)", konwertuj_do_docx(res.text), "Protokol_MeatPoint.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    except Exception as e: 
                        st.error(czytelny_blad(e, " podczas generowania opisu"))

# ==============================================================================
# 🎙️ ZAKŁADKA 2: GŁOSOWY EDYTOR PROTOKOŁÓW — RÓWNOLEGŁE PRZETWARZANIE
# ==============================================================================
with tab2:
    st.title("🎙️ Edytor głosowy opisów wizyt")
    
    if 'sekcje_dokumentu' not in st.session_state: st.session_state.sekcje_dokumentu = None
    if 'koszyk_nagran' not in st.session_state: st.session_state.koszyk_nagran = {}
    if 'koszyk_tekstowy' not in st.session_state: st.session_state.koszyk_tekstowy = {}
    if 'v_key' not in st.session_state: st.session_state.v_key = str(uuid.uuid4())
    if 'klucze_mikrofonow' not in st.session_state: st.session_state.klucze_mikrofonow = {}
    if 'wynik_poprawek' not in st.session_state: st.session_state.wynik_poprawek = None

    col_top1, col_top2 = st.columns([3, 1])
    with col_top1:
        u_file = st.file_uploader("📂 Wgraj plik opisu (.docx):", type=["docx"], key=f"u_{st.session_state.v_key}")
    with col_top2:
        st.write("<br/>", unsafe_allow_html=True)
        if st.button("🔄 Nowy protokół / Reset", type="secondary", use_container_width=True):
            st.session_state.sekcje_dokumentu = None
            st.session_state.koszyk_nagran = {}
            st.session_state.koszyk_tekstowy = {}
            st.session_state.klucze_mikrofonow = {}
            st.session_state.v_key = str(uuid.uuid4())
            st.session_state.wynik_poprawek = None
            st.rerun()
            
    if u_file and st.session_state.sekcje_dokumentu is None:
        if st.button("⚙️ Załaduj strukturę pliku"):
            st.session_state.sekcje_dokumentu = segmentuj_docx(u_file.read()); st.rerun()

    if st.session_state.sekcje_dokumentu:
        # Trwały komunikat po wprowadzeniu poprawek — przetrwa przeładowanie strony (st.rerun)
        if st.session_state.wynik_poprawek:
            wynik = st.session_state.wynik_poprawek
            if wynik["bledy"]:
                for b in wynik["bledy"]:
                    st.error(b)
                if wynik["ok"]:
                    st.warning(f"⚠️ Wprowadzono {wynik['ok']} poprawek, {len(wynik['bledy'])} nie udało się przetworzyć.")
            else:
                st.success(f"✅ Gotowe! Wszystkie {wynik['ok']} poprawki naniesione na protokół.")
            if wynik["ok"]:
                st.info("⬇️ Zjedź na dół do sekcji **3️⃣ Pobieranie gotowego dokumentu** i kliknij **📦 Generuj finalny plik Word z poprawkami**, aby pobrać zaktualizowany protokół.")
            st.session_state.wynik_poprawek = None  # pokaż raz, potem wyczyść

        st.markdown("---")
        col_ed1, col_ed2 = st.columns([1, 1], gap="large")
        
        with col_ed1:
            st.markdown("### 1️⃣ Wybór obszaru do korekty")
            wybrana_sekcja = st.selectbox("Wybierz nagłówek, do którego chcesz dodać nagranie:", list(st.session_state.sekcje_dokumentu.keys()), key="sel_voice", format_func=czytelna_nazwa)

            # Podgląd treści: STAŁY klucz + wartość podawana przez pamięć sesji.
            # Dzięki temu element nie jest przebudowywany przy zmianie sekcji (brak przeskoku strony).
            st.session_state["podglad_biezacej_sekcji"] = st.session_state.sekcje_dokumentu[wybrana_sekcja]
            st.text_area("📄 Aktualna treść sekcji:", height=180, disabled=True, key="podglad_biezacej_sekcji")
            
            if wybrana_sekcja not in st.session_state.klucze_mikrofonow:
                st.session_state.klucze_mikrofonow[wybrana_sekcja] = str(uuid.uuid4())
            
            mic_id = f"mic_{wybrana_sekcja}_{st.session_state.klucze_mikrofonow[wybrana_sekcja]}"
            audio_instrukcja = mic_recorder(start_prompt="🎙️ Nagraj uwagę dla tej sekcji", stop_prompt="🛑 Zatrzymaj i zapisz w pamięci", key=mic_id)
            
            if audio_instrukcja:
                if wybrana_sekcja not in st.session_state.koszyk_nagran or st.session_state.koszyk_nagran[wybrana_sekcja] != audio_instrukcja['bytes']:
                    st.session_state.koszyk_nagran[wybrana_sekcja] = audio_instrukcja['bytes']
                    st.rerun()

            st.markdown("**✏️ Lub wpisz uwagę tekstową:**")
            # Przy zmianie sekcji czyścimy pole (stały klucz = brak przeskoku strony,
            # a pole i tak startuje puste dla każdej nowej sekcji).
            if st.session_state.get("ostatnia_sekcja_edytora") != wybrana_sekcja:
                st.session_state["uwaga_tekstowa_input"] = ""
                st.session_state["ostatnia_sekcja_edytora"] = wybrana_sekcja
            uwaga_tekstowa = st.text_area(
                "Wpisz krótkie zalecenia / korektę dla tej sekcji:",
                height=100,
                key="uwaga_tekstowa_input",
                placeholder="Np.: waga 4.2 kg, podawać 2x dziennie, unikać ryby"
            )
            if st.button("💾 Zapisz uwagę tekstową", key="btn_zapisz_uwage", use_container_width=True):
                tekst = uwaga_tekstowa.strip()
                if tekst:
                    st.session_state.koszyk_tekstowy[wybrana_sekcja] = tekst
                    st.session_state["ostatnia_sekcja_edytora"] = None  # wymuś wyczyszczenie pola po zapisie
                    st.toast(f"✅ Zapisano uwagę tekstową dla: {czytelna_nazwa(wybrana_sekcja)}")
                    st.rerun()
                else:
                    st.warning("⚠️ Pole uwagi jest puste.")

        with col_ed2:
            st.markdown("### 2️⃣ Kolejka uwag do wprowadzenia")

            jest_cos = st.session_state.koszyk_nagran or st.session_state.koszyk_tekstowy

            if not jest_cos:
                st.info("Brak oczekujących uwag. Wybierz sekcję po lewej stronie, nagraj głos lub wpisz uwagę tekstową.")
            else:
                # --- Uwagi głosowe ---
                if st.session_state.koszyk_nagran:
                    st.markdown("**🎙️ Uwagi głosowe:**")
                    for s_nazwa in list(st.session_state.koszyk_nagran.keys()):
                        if s_nazwa in st.session_state.koszyk_nagran:
                            a_bytes = st.session_state.koszyk_nagran[s_nazwa]
                            c_box1, c_box2 = st.columns([5, 1])
                            c_box1.markdown(f"**📌 {czytelna_nazwa(s_nazwa)}**")
                            c_box1.audio(a_bytes, format="audio/wav")
                            if c_box2.button("❌", key=f"del_glos_{s_nazwa}", help="Usuń to nagranie"):
                                del st.session_state.koszyk_nagran[s_nazwa]
                                st.session_state.klucze_mikrofonow[s_nazwa] = str(uuid.uuid4())
                                st.toast(f"🗑️ Usunięto nagranie głosowe: {czytelna_nazwa(s_nazwa)}")
                                st.rerun()

                # --- Uwagi tekstowe ---
                if st.session_state.koszyk_tekstowy:
                    st.markdown("**✏️ Uwagi tekstowe:**")
                    for s_nazwa in list(st.session_state.koszyk_tekstowy.keys()):
                        if s_nazwa in st.session_state.koszyk_tekstowy:
                            t_tresc = st.session_state.koszyk_tekstowy[s_nazwa]
                            c_box1, c_box2 = st.columns([5, 1])
                            c_box1.markdown(f"**📌 {czytelna_nazwa(s_nazwa)}**")
                            c_box1.caption(t_tresc)
                            if c_box2.button("❌", key=f"del_txt_{s_nazwa}", help="Usuń tę uwagę tekstową"):
                                del st.session_state.koszyk_tekstowy[s_nazwa]
                                st.toast(f"🗑️ Usunięto uwagę tekstową: {czytelna_nazwa(s_nazwa)}")
                                st.rerun()

                st.markdown("---")
                liczba_wszystkich = len(st.session_state.koszyk_nagran) + len(st.session_state.koszyk_tekstowy)
                if st.button(f"🚀 WPROWADŹ WSZYSTKIE POPRAWKI ({liczba_wszystkich} szt.)", type="primary", use_container_width=True):
                    if not api_key: st.error("❌ Podaj klucz API Gemini!")
                    else:
                        client = genai.Client(api_key=api_key)
                        progress_bar = st.progress(0, text=f"Przygotowanie... (0 / {liczba_wszystkich})")
                        wyniki = {}
                        bledy = []
                        ukonczone = 0

                        # --- BLOK 1: Równoległe przetwarzanie nagrań głosowych ---
                        if st.session_state.koszyk_nagran:
                            zadania_glos = [
                                (
                                    s_nazwa,
                                    a_bytes,
                                    st.session_state.sekcje_dokumentu.get(s_nazwa, ""),
                                    api_key,
                                    model_choice,
                                )
                                for s_nazwa, a_bytes in list(st.session_state.koszyk_nagran.items())
                            ]
                            try:
                                with ThreadPoolExecutor(max_workers=min(10, len(zadania_glos))) as executor:
                                    futures = {
                                        executor.submit(przetworz_jedno_nagranie, z): z[0]
                                        for z in zadania_glos
                                    }
                                    for future in as_completed(futures):
                                        s_nazwa = futures[future]
                                        ukonczone += 1
                                        try:
                                            nazwa, nowa_tresc = future.result()
                                            wyniki[nazwa] = nowa_tresc
                                            progress_bar.progress(
                                                ukonczone / liczba_wszystkich,
                                                text=f"Przetworzono: {ukonczone} / {liczba_wszystkich} — ✅ 🎙️ {czytelna_nazwa(nazwa)}"
                                            )
                                        except Exception as e:
                                            bledy.append(
                                                f"🎙️ **Sekcja: {czytelna_nazwa(s_nazwa)} (uwaga głosowa)**\n\n"
                                                + czytelny_blad(e)
                                            )
                                            progress_bar.progress(
                                                ukonczone / liczba_wszystkich,
                                                text=f"Przetworzono: {ukonczone} / {liczba_wszystkich} — ⚠️ {czytelna_nazwa(s_nazwa)}"
                                            )
                            except Exception as e:
                                bledy.append(czytelny_blad(e, " podczas przetwarzania nagrań"))

                        # --- BLOK 2: Sekwencyjne przetwarzanie uwag tekstowych ---
                        # (szybkie - tylko redakcja tekstu, bez audio, nie wymaga wątków)
                        if st.session_state.koszyk_tekstowy:
                            for s_nazwa, uwaga in list(st.session_state.koszyk_tekstowy.items()):
                                ukonczone += 1
                                try:
                                    # Jeśli ta sekcja była już zmieniona przez głos, bierzemy nową treść
                                    baza = wyniki.get(s_nazwa, st.session_state.sekcje_dokumentu.get(s_nazwa, ""))
                                    p_txt = (
                                        f"Jesteś redaktorem dokumentacji medycznej. Masz dwa teksty:\n\n"
                                        f"ORYGINALNY TEKST SEKCJI '{s_nazwa}':\n{baza}\n\n"
                                        f"UWAGA DO WPLECENIA (napisana przez dietetyka w skrócie):\n{uwaga}\n\n"
                                        f"🚨 BEZWZGLĘDNE ZASADY — CZYTAJ UWAŻNIE:\n"
                                        f"1. Wpleć treść UWAGI DO WPLECENIA w ORYGINALNY TEKST SEKCJI.\n"
                                        f"2. ZAKAZ dodawania JAKICHKOLWIEK słów, faktów, zaleceń, uzupełnień, których NIE MA w obu powyższych tekstach. Zero inwencji własnej.\n"
                                        f"3. Twoim jedynym zadaniem jest ułożenie podanych informacji w poprawny gramatycznie i spójny tekst po polsku.\n"
                                        f"4. ZAKAZ wstępów, podsumowań, komentarzy AI ('Oto zaktualizowana sekcja:' itp.).\n"
                                        f"5. Zwróć WYŁĄCZNIE gotowy, czysty tekst sekcji."
                                    )
                                    wynik_txt = generuj_z_retry(client, model_choice, p_txt).text.strip()
                                    wyniki[s_nazwa] = wynik_txt
                                    progress_bar.progress(
                                        ukonczone / liczba_wszystkich,
                                        text=f"Przetworzono: {ukonczone} / {liczba_wszystkich} — ✅ ✏️ {czytelna_nazwa(s_nazwa)}"
                                    )
                                except Exception as e:
                                    bledy.append(
                                        f"✏️ **Sekcja: {czytelna_nazwa(s_nazwa)} (uwaga tekstowa)**\n\n"
                                        + czytelny_blad(e)
                                    )
                                    progress_bar.progress(
                                        ukonczone / liczba_wszystkich,
                                        text=f"Przetworzono: {ukonczone} / {liczba_wszystkich} — ⚠️ {czytelna_nazwa(s_nazwa)}"
                                    )

                        # --- Zastosuj wyniki i wyczyść koszyki ---
                        for nazwa, nowa_tresc in wyniki.items():
                            st.session_state.sekcje_dokumentu[nazwa] = nowa_tresc
                            if nazwa in st.session_state.klucze_mikrofonow:
                                st.session_state.klucze_mikrofonow[nazwa] = str(uuid.uuid4())
                        for nazwa in wyniki:
                            st.session_state.koszyk_nagran.pop(nazwa, None)
                            st.session_state.koszyk_tekstowy.pop(nazwa, None)

                        progress_bar.empty()

                        # Zapisz wynik do pamięci sesji — komunikat pokaże się PO przeładowaniu,
                        # na górze edytora, gdzie jest trwały i widoczny dla użytkownika.
                        st.session_state.wynik_poprawek = {"ok": len(wyniki), "bledy": list(bledy)}

                        st.rerun()

        st.markdown("---")
        st.markdown("### 3️⃣ Pobieranie gotowego dokumentu")
        if st.button("📦 Generuj finalny plik Word z poprawkami", type="primary", key="btn_build_final"):
            t_md = ""
            for sk, ts in st.session_state.sekcje_dokumentu.items():
                if sk in ["Nagłówek i Metryczka", "Nagłówek i Data wizyty"]: t_md += f"{ts}\n\n"
                else:
                    prefix = "" if sk.startswith("###") else "## "
                    t_md += f"{prefix}{sk}\n{ts}\n\n"
            st.download_button("📥 POBIERZ PROTOKÓŁ (.DOCX)", konwertuj_do_docx(t_md), "Protokol_MeatPoint_Poprawiony.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
